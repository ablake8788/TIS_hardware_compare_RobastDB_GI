"""
Renderer: converts AnalysisResult into a polished .docx Word report.
Uses python-docx. No workflow logic — formatting/conversion only.
"""
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from domain.models import AnalysisResult, ComparisonRow

# ── Color palette ────────────────────────────────────────────────────────────
GREEN_FILL   = "C6EFCE"   # Yes
YELLOW_FILL  = "FFEB9C"   # Partial
RED_FILL     = "FFC7CE"   # No
HEADER_FILL  = "1F4E79"   # Table header (dark blue)
HEADER_TXT   = "FFFFFF"
ALT_ROW      = "EBF3FB"   # Alternating row tint
TITLE_COLOR  = "1F4E79"
ACCENT_COLOR = "2E75B6"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _para_in_cell(cell, text: str, bold=False, font_size=9,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return para


def _add_section_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(4)
    return p


def render_docx(result: AnalysisResult, output_path: str) -> str:
    """Render AnalysisResult to a .docx file. Returns the file path."""
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # ── Default font ─────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    t_run = title.add_run("Hardware Gap Analysis Report")
    t_run.bold = True
    t_run.font.size = Pt(20)
    t_run.font.color.rgb = RGBColor.from_string(TITLE_COLOR)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run(
        f"Titanium  vs  {result.competitor.name}  ·  "
        f"Generated {result.generated_at.strftime('%B %d, %Y %H:%M')}"
    )
    s_run.font.size = Pt(11)
    s_run.font.color.rgb = RGBColor.from_string(ACCENT_COLOR)
    subtitle.paragraph_format.space_after = Pt(16)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(0)
    hr_pPr = hr._p.get_or_add_pPr()
    hr_pBdr = OxmlElement("w:pBdr")
    hr_bottom = OxmlElement("w:bottom")
    hr_bottom.set(qn("w:val"), "single")
    hr_bottom.set(qn("w:sz"), "6")
    hr_bottom.set(qn("w:space"), "1")
    hr_bottom.set(qn("w:color"), ACCENT_COLOR)
    hr_pBdr.append(hr_bottom)
    hr_pPr.append(hr_pBdr)

    # ── Executive summary cards ───────────────────────────────────────────────
    doc.add_paragraph()
    _add_section_heading(doc, "Executive Summary")

    summary_tbl = doc.add_table(rows=2, cols=4)
    summary_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_w = Inches(1.6)
    labels = ["Total Hardware", f"{result.competitor.name} — Yes",
              f"{result.competitor.name} — Partial", f"{result.competitor.name} — No"]
    values = [str(result.total), str(result.yes_count),
              str(result.partial_count), str(result.no_count)]
    fills  = [HEADER_FILL, GREEN_FILL, YELLOW_FILL, RED_FILL]
    txt_colors = [HEADER_TXT, "375623", "7D4A00", "9C0006"]

    for i, (lbl, val, fill, tc) in enumerate(zip(labels, values, fills, txt_colors)):
        hdr_cell = summary_tbl.rows[0].cells[i]
        val_cell = summary_tbl.rows[1].cells[i]
        hdr_cell.width = col_w
        val_cell.width = col_w
        _set_cell_bg(hdr_cell, fill)
        _set_cell_bg(val_cell, fill)
        _set_cell_borders(hdr_cell)
        _set_cell_borders(val_cell)
        _para_in_cell(hdr_cell, lbl, bold=True, font_size=9, color=HEADER_TXT if fill == HEADER_FILL else tc)
        _para_in_cell(val_cell, val, bold=True, font_size=22, color=tc, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()

    # ── Legend ────────────────────────────────────────────────────────────────
    legend_para = doc.add_paragraph()
    legend_para.paragraph_format.space_before = Pt(4)
    legend_para.paragraph_format.space_after = Pt(8)
    for label, color in [("  Yes = Full match  ", GREEN_FILL),
                          ("  Partial = Limited/bundled/integrated  ", YELLOW_FILL),
                          ("  No = Not offered  ", RED_FILL)]:
        run = legend_para.add_run(label)
        run.font.size = Pt(8)
        run.font.bold = True

    # ── Comparison table ──────────────────────────────────────────────────────
    _add_section_heading(doc, "Hardware Comparison Matrix")

    headers = ["Hardware", "Description", "Companies Using", "Titanium",
               result.competitor.name, "Notes"]
    col_widths = [Inches(1.4), Inches(2.0), Inches(1.4), Inches(0.7), Inches(0.85), Inches(1.35)]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, (hdr, w) in enumerate(zip(headers, col_widths)):
        cell = hdr_row.cells[i]
        cell.width = w
        _set_cell_bg(cell, HEADER_FILL)
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _para_in_cell(cell, hdr, bold=True, font_size=9, color=HEADER_TXT)

    # Data rows
    for idx, row in enumerate(result.rows):
        tr = table.add_row()
        use_alt = (idx % 2 == 1)
        base_fill = ALT_ROW if use_alt else "FFFFFF"

        values = [
            row.hardware, row.description, row.companies_using,
            row.titanium, row.competitor, row.competitor_notes
        ]
        fills_row = [
            base_fill, base_fill, base_fill,
            (GREEN_FILL if row.titanium == "Yes" else RED_FILL),
            (GREEN_FILL if row.competitor == "Yes"
             else YELLOW_FILL if row.competitor == "Partial"
             else RED_FILL),
            base_fill
        ]
        text_colors_row = [
            "000000", "444444", "444444",
            ("375623" if row.titanium == "Yes" else "9C0006"),
            ("375623" if row.competitor == "Yes"
             else "7D4A00" if row.competitor == "Partial"
             else "9C0006"),
            "555555"
        ]
        bolds = [True, False, False, True, True, False]

        for i, (val, fill, tc, bold, w) in enumerate(
                zip(values, fills_row, text_colors_row, bolds, col_widths)):
            cell = tr.cells[i]
            cell.width = w
            _set_cell_bg(cell, fill)
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _para_in_cell(cell, val, bold=bold, font_size=9, color=tc)

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run(
        f"⚠  AI-generated analysis — verify against official {result.competitor.name} "
        "product documentation before use in sales or procurement decisions."
    )
    note_run.italic = True
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor.from_string("888888")

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else ".", exist_ok=True)
    doc.save(output_path)
    return output_path
